import os
import re
import uuid
import json
import asyncio
import sys
from typing import Dict, List, Any, Optional, Union, Tuple
from datetime import datetime, timedelta
from pathlib import Path
from enum import Enum
import logging
from dataclasses import dataclass, asdict

# Core LangChain imports
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.memory import ConversationSummaryBufferMemory
from langchain.schema import Document

# Enhanced document processing
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import pymupdf  # PyMuPDF for better PDF handling
from docx import Document as DocxDocument
import io
import base64
from concurrent.futures import ThreadPoolExecutor

# API Key
GOOGLE_API_KEY = "AIzaSyATbLAXPhXEwRJeV6Svajr8pX8AWfNiZCQ"

# Minimal logging
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

class ContentType(Enum):
    TEXT = "text"
    DIAGRAM = "diagram" 
    HANDWRITTEN = "handwritten"
    TABLE = "table"
    MIXED = "mixed"

class TeachingStyle(Enum):
    FORMAL = "formal"
    ANALOGY = "analogy"
    BALANCED = "balanced"
    CONCISE = "concise"

@dataclass
class ProcessingResult:
    content: str
    content_type: ContentType
    confidence: float
    metadata: Dict[str, Any]
    processing_time: float

class EnhancedStudyAssistant:
    """24/7 Personalized Academic Tutor - Your Study Companion"""
    
    def __init__(self, memory_manager=None, max_tokens: int = 2000):
        self.agent_name = "Enhanced Study Assistant"
        self.memory_manager = memory_manager or self._create_default_memory()
        self.max_tokens = max_tokens
        
        # Initialize models
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=GOOGLE_API_KEY,
            temperature=0.4,
            max_output_tokens=1500,
        )
        
        self.vision_llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-pro",
            google_api_key=GOOGLE_API_KEY,
            temperature=0.1,
            max_output_tokens=1000,
        )
        
        # Enhanced memory with summarization
        self.memory = ConversationSummaryBufferMemory(
            llm=self.llm,
            max_token_limit=self.max_tokens,
            return_messages=True,
            memory_key="chat_history"
        )
        
        # Track user preferences
        self.user_preferences = {
            "teaching_style": TeachingStyle.BALANCED,
            "detail_level": "moderate",
            "question_frequency": "balanced"
        }
        
        # Thread pool for async file processing
        self.executor = ThreadPoolExecutor(max_workers=3)
        
        # Setup prompts
        self._setup_prompts()
        
    def _create_default_memory(self):
        """Create default memory manager if none provided"""
        class DefaultMemoryManager:
            def __init__(self):
                self.shared_memory = {
                    'user_profile': {
                        'major': 'General Studies',
                        'learning_style': 'Visual',
                        'level': 'Undergraduate'
                    },
                    'academic_data': {
                        'recent_topics': [],
                        'weak_subjects': []
                    }
                }
            
            def add_to_list(self, category, key, value):
                if category not in self.shared_memory:
                    self.shared_memory[category] = {}
                if key not in self.shared_memory[category]:
                    self.shared_memory[category][key] = []
                if isinstance(value, dict):
                    self.shared_memory[category][key].append(value)
                elif value not in self.shared_memory[category][key]:
                    self.shared_memory[category][key].append(value)
        
        return DefaultMemoryManager()
        
    def _setup_prompts(self):
        """Setup structured prompts for tutoring"""
        
        self.conversation_prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert academic tutor and study companion. 

# TEACHING APPROACH ADJUSTMENTS:
1. **ANSWER FIRST**: Provide the main answer or explanation upfront, then ask follow-up questions
2. **ADAPT TO USER PREFERENCE**: 
   - If user asks for formal explanation: provide structured, detailed answers
   - If user asks for analogy: use relatable comparisons
   - If user seems impatient: be concise and direct
3. **BALANCE QUESTIONS**: Use questions to check understanding, but don't overdo it
4. **BE DIRECT**: When the question is clear, provide a clear answer first

# RESPONSE GUIDELINES:
- Start with the core answer/explanation
- Then ask 1-2 thoughtful questions to check understanding
- Use formatting (bullets, numbering) for clarity
- Adjust formality based on user cues
- If user uploads content, analyze it first then respond

Current user preferences: {teaching_style}
Detail level: {detail_level}"""),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}")
        ])
        
    def _detect_user_preferences(self, user_input: str) -> Dict[str, Any]:
        """Detect user preferences from their input"""
        input_lower = user_input.lower()
        
        # Detect teaching style preference
        if any(word in input_lower for word in ["formal", "detailed", "technical", "academic"]):
            teaching_style = TeachingStyle.FORMAL
        elif any(word in input_lower for word in ["analogy", "example", "like a", "similar to", "metaphor"]):
            teaching_style = TeachingStyle.ANALOGY
        elif any(word in input_lower for word in ["concise", "brief", "short", "quick"]):
            teaching_style = TeachingStyle.CONCISE
        else:
            teaching_style = self.user_preferences["teaching_style"]
        
        # Detect detail level preference
        if any(word in input_lower for word in ["detailed", "in-depth", "comprehensive", "thorough"]):
            detail_level = "high"
        elif any(word in input_lower for word in ["summary", "overview", "gist", "main points"]):
            detail_level = "low"
        else:
            detail_level = self.user_preferences["detail_level"]
            
        # Update preferences
        self.user_preferences.update({
            "teaching_style": teaching_style,
            "detail_level": detail_level
        })
        
        return self.user_preferences
    
    def _format_response_based_on_preferences(self, response: str, preferences: Dict[str, Any]) -> str:
        """Format the response based on user preferences"""
        teaching_style = preferences["teaching_style"]
        detail_level = preferences["detail_level"]
        
        if teaching_style == TeachingStyle.FORMAL:
            # Make response more structured and academic
            response = self._make_response_formal(response)
        elif teaching_style == TeachingStyle.ANALOGY:
            # Add analogies if not already present
            if "like" not in response.lower() and "similar to" not in response.lower():
                response += "\n\n💡 Think of it like: [add your own analogy based on the context]"
        elif teaching_style == TeachingStyle.CONCISE:
            # Make response more concise
            response = self._make_response_concise(response)
        
        # Adjust detail level
        if detail_level == "low" and len(response.split()) > 150:
            response = self._summarize_response(response)
        elif detail_level == "high" and len(response.split()) < 100:
            response = self._expand_response(response)
            
        return response
    
    def _make_response_formal(self, response: str) -> str:
        """Make response more formal and structured"""
        # Add academic structure if needed
        if not response.startswith(("Firstly", "In summary", "There are")):
            # Try to add some structure
            lines = response.split('\n')
            if len(lines) > 2:
                response = f"In addressing this question:\n\n1. {lines[0]}\n2. {lines[1]}\n3. {lines[2]}" + \
                          ("\n\n" + '\n'.join(lines[3:]) if len(lines) > 3 else "")
        return response
    
    def _make_response_concise(self, response: str) -> str:
        """Make response more concise"""
        # Try to shorten the response while keeping key information
        sentences = response.split('. ')
        if len(sentences) > 3:
            response = '. '.join(sentences[:3]) + '.'
            if "In short" not in response and "Basically" not in response:
                response = "In short: " + response
        return response
    
    def _summarize_response(self, response: str) -> str:
        """Summarize a lengthy response"""
        # Simple summarization - in a real implementation, you'd use more sophisticated methods
        sentences = response.split('. ')
        if len(sentences) > 4:
            key_sentences = [sentences[0]] + [s for s in sentences if any(word in s.lower() 
                                                                        for word in ['important', 'key', 'main', 'crucial'])]
            if len(key_sentences) < 2 and len(sentences) > 2:
                key_sentences.append(sentences[1])
            if len(key_sentences) < 3 and len(sentences) > 3:
                key_sentences.append(sentences[-1])
            response = '. '.join(key_sentences[:4]) + '.'
        return response
    
    def _expand_response(self, response: str) -> str:
        """Expand a brief response"""
        # Simple expansion - in a real implementation, you'd use more sophisticated methods
        if len(response.split()) < 80:
            if "For example" not in response and "Specifically" not in response:
                response += " For example, [add a specific example related to the topic]."
            if "This is important because" not in response:
                response += " This is important because [explain relevance]."
        return response
    
    async def chat(self, user_input: str, file_data: Optional[bytes] = None, 
                   file_type: Optional[str] = None, filename: Optional[str] = None) -> str:
        """Main chat method for conversational interaction"""
        
        try:
            # Detect user preferences
            preferences = self._detect_user_preferences(user_input)
            
            # Get contexts
            user_context = self._get_user_context()
            session_context = self._get_session_context()
            
            processed_content = ""
            
            # Process file if provided
            if file_data and file_type:
                processing_result = await self._process_file_async(file_data, file_type, filename)
                processed_content = self._format_processing_result(processing_result)
            
            # Prepare conversation chain
            chain = (
                RunnablePassthrough.assign(
                    user_context=lambda _: user_context,
                    session_context=lambda _: session_context,
                    teaching_style=lambda _: preferences["teaching_style"].value,
                    detail_level=lambda _: preferences["detail_level"],
                    chat_history=lambda _: self.memory.chat_memory.messages
                )
                | self.conversation_prompt
                | self.llm
                | StrOutputParser()
            )
            
            # Combine processed content with user input
            full_input = f"{processed_content}\n\nStudent Question: {user_input}" if processed_content else user_input
            
            # Generate response
            response = await chain.ainvoke({"input": full_input})
            
            # Format based on preferences
            formatted_response = self._format_response_based_on_preferences(response, preferences)
            
            # Update memory
            self.memory.save_context(
                {"input": user_input}, 
                {"output": formatted_response}
            )
            
            # Track session
            await self._track_study_session(user_input, formatted_response, file_type, filename)
            
            return formatted_response
            
        except Exception as e:
            return f"I apologize, but I encountered an issue: {str(e)}. Could you please try rephrasing your question?"
    
    async def _process_file_async(self, file_data: bytes, file_type: str, filename: str) -> ProcessingResult:
        """Process uploaded files"""
        start_time = datetime.now()
        
        try:
            if file_type.startswith('image'):
                result = await self._process_image_advanced(file_data)
            elif file_type == 'application/pdf':
                result = await self._process_pdf_advanced(file_data)
            elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                result = await self._process_docx_advanced(file_data)
            else:
                result = ProcessingResult(
                    content=f"I can see you uploaded a {file_type} file. I work best with images, PDFs, and Word documents.",
                    content_type=ContentType.TEXT,
                    confidence=0.5,
                    metadata={"file_type": file_type},
                    processing_time=0.0
                )
            
            result.processing_time = (datetime.now() - start_time).total_seconds()
            return result
            
        except Exception as e:
            return ProcessingResult(
                content=f"I had trouble processing your file: {str(e)}",
                content_type=ContentType.TEXT,
                confidence=0.0,
                metadata={"error": str(e)},
                processing_time=(datetime.now() - start_time).total_seconds()
            )
    
    async def _process_image_advanced(self, image_data: bytes) -> ProcessingResult:
        """Process images with OCR and enhancement"""
        start_time = datetime.now()
        try:
            image = Image.open(io.BytesIO(image_data))
            text = pytesseract.image_to_string(image)

            if text.strip():
                return ProcessingResult(
                    content=f"Text extracted via OCR:\n{text}",
                    content_type=ContentType.TEXT,
                    confidence=0.85,
                    metadata={"method": "ocr"},
                    processing_time=(datetime.now()-start_time).total_seconds()
                )
            else:
                return ProcessingResult(
                    content="No text found in image via OCR.",
                    content_type=ContentType.MIXED,
                    confidence=0.5,
                    metadata={"method": "ocr"},
                    processing_time=(datetime.now()-start_time).total_seconds()
                )
        except Exception as e:
            return ProcessingResult(
                content=f"Image processing failed: {str(e)}",
                content_type=ContentType.TEXT,
                confidence=0.0,
                metadata={"error": str(e)},
                processing_time=(datetime.now()-start_time).total_seconds()
            )
    
    async def _process_pdf_advanced(self, pdf_data: bytes) -> ProcessingResult:
        """Process PDF documents"""
        try:
            pdf_document = pymupdf.open(stream=pdf_data, filetype="pdf")
            text_content = []
            
            for page_num in range(min(5, len(pdf_document))):  # Limit to first 5 pages
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    text_content.append(f"Page {page_num + 1}:\n{text[:1000]}...")  # Limit text length
            
            pdf_document.close()
            
            return ProcessingResult(
                content="\n\n".join(text_content),
                content_type=ContentType.TEXT,
                confidence=0.9,
                metadata={"pages": len(pdf_document)}
            )
            
        except Exception as e:
            raise Exception(f"PDF processing failed: {str(e)}")
    
    async def _process_docx_advanced(self, doc_data: bytes) -> ProcessingResult:
        """Process Word documents"""
        try:
            doc = DocxDocument(io.BytesIO(doc_data))
            content_parts = []
            
            for para in doc.paragraphs[:20]:  # Limit to first 20 paragraphs
                if para.text.strip():
                    content_parts.append(para.text)
            
            return ProcessingResult(
                content="\n\n".join(content_parts),
                content_type=ContentType.TEXT,
                confidence=0.95,
                metadata={"paragraphs": len(doc.paragraphs)}
            )
            
        except Exception as e:
            raise Exception(f"Word processing failed: {str(e)}")
    
    def _format_processing_result(self, result: ProcessingResult) -> str:
        """Format file processing results"""
        confidence_emoji = "🟢" if result.confidence > 0.8 else "🟡" if result.confidence > 0.6 else "🔴"
        
        return f"""
📄 **File Analysis Complete** {confidence_emoji}

{result.content}

---
"""
    
    def _get_user_context(self) -> str:
        """Get user context for personalization"""
        try:
            profile = self.memory_manager.shared_memory.get('user_profile', {})
            academic = self.memory_manager.shared_memory.get('academic_data', {})
            
            context_parts = [
                f"Major: {profile.get('major', 'General Studies')}",
                f"Learning Style: {profile.get('learning_style', 'Visual')}",
                f"Level: {profile.get('level', 'Undergraduate')}"
            ]
            
            recent_topics = academic.get('recent_topics', [])[-3:]
            if recent_topics:
                context_parts.append(f"Recent Topics: {', '.join(recent_topics)}")
            
            return "\n".join(context_parts)
        except:
            return "New student profile"
    
    def _get_session_context(self) -> str:
        """Get current session context"""
        try:
            messages = self.memory.chat_memory.messages[-2:]
            if messages:
                return "Continuing previous conversation"
            return "Starting new conversation"
        except:
            return "New session"
    
    async def _track_study_session(self, user_input: str, response: str, 
                                 file_type: Optional[str], filename: Optional[str]):
        """Track study sessions for analytics"""
        try:
            # Simple subject detection
            subjects = {
                'mathematics': ['math', 'calculus', 'algebra', 'geometry', 'statistics'],
                'physics': ['physics', 'force', 'energy', 'velocity', 'quantum'],
                'chemistry': ['chemistry', 'element', 'compound', 'reaction', 'organic'],
                'biology': ['biology', 'cell', 'dna', 'organism', 'evolution'],
                'computer_science': ['programming', 'algorithm', 'code', 'software'],
                'literature': ['literature', 'poem', 'novel', 'author', 'writing'],
                'history': ['history', 'war', 'revolution', 'historical']
            }
            
            combined_text = f"{user_input} {response}".lower()
            detected_subjects = []
            
            for subject, keywords in subjects.items():
                if any(keyword in combined_text for keyword in keywords):
                    detected_subjects.append(subject)
            
            # Update memory manager
            for subject in detected_subjects:
                self.memory_manager.add_to_list('academic_data', 'recent_topics', subject)
                
        except Exception:
            pass  # Don't let tracking errors break the conversation

# Simple conversation loop
async def start_conversation(user_input):
    """Start the study assistant conversation"""
    
    print("🎓 Enhanced Study Assistant - Your Personal Tutor")
    print("=" * 50)
    print("Hi! I'm here to help you learn and understand any topic.")
    print("You can ask me questions, upload images/documents, or request study materials.")
    print("Type 'exit' or 'quit' to end our session.\n")
    
    # Initialize assistant
    assistant = EnhancedStudyAssistant()
    
    while True:
        try: 
            if user_input.lower() in ['exit', 'quit', 'bye']:
                return "bravo! for today!"
                break
            
            if not user_input:
                continue
            
            # Get response from assistant
            print("\n🤔 Thinking...")
            response = await assistant.chat(user_input)
            
            # Display response
            print(f"\n📚 Tutor: {response}\n")
            print("-" * 50)
            return response
            
        except KeyboardInterrupt:
            print("\n\n👋 Study session ended. Good luck with your studies!")
            break
        except Exception as e:
            print(f"\n❌ Oops, something went wrong: {e}")
            print("Let's try again!\n")

            

# Add this at the bottom, before the "__main__" block:

async def test_file(file_path: str):
    """Test the agent's PDF, Word, or image processing"""
    assistant = EnhancedStudyAssistant()
    
    # Detect file type
    ext = os.path.splitext(file_path)[1].lower()
    if ext in ['.pdf']:
        file_type = 'application/pdf'
    elif ext in ['.docx']:
        file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif ext in ['.png', '.jpg', '.jpeg']:
        file_type = f'image/{ext[1:]}'
    else:
        print(f"Unsupported file type: {ext}")
        return
    
    # Read file
    with open(file_path, 'rb') as f:
        file_bytes = f.read()
    
    print(f"\n🧪 Testing file: {file_path}\n")
    response = await assistant.chat(
        user_input="Please analyze this file.",
        file_data=file_bytes,
        file_type=file_type,
        filename=os.path.basename(file_path)
    )
    
    print(f"📚 Agent Response:\n{response}\n")
    print("="*50)


if __name__ == "__main__":
    # Check if a file path is passed as argument
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            asyncio.run(test_file(file_path))
        except Exception as e:
            print(f"Error testing file: {e}")
    else:
        # Run interactive conversation
        print("Starting your personal study assistant...")
        try:
            asyncio.run(start_conversation())
        except KeyboardInterrupt:
            print("\nGoodbye!")
        except Exception as e:
            print(f"Error: {e}")